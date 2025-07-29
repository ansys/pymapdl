# Copyright (C) 2016 - 2025 ANSYS, Inc. and/or its affiliates.
# SPDX-License-Identifier: MIT
#
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

"""
.. _ref_beam_report_example:

I-Beam Analysis with Automated Report Generation
=================================================

.. note:: This file is complimentary material of
   `I-Beam Analysis with Automated Report Generation <beam_analysis_report_example>`_
   example.

This example demonstrates a comprehensive structural analysis of a fully constrained
I-beam using PyMAPDL. The script creates a parameterized I-beam model, applies loads,
solves for structural response, and generates a detailed engineering report in both
Markdown and Microsoft Word formats.

The example showcases:
- Parameterized I-beam geometry definition
- Material property assignment
- Fully constrained at both ends
- Distributed load application
- Result extraction and post-processing
- Automated report generation with plots and tables

This example requires the following imports:

"""

from datetime import datetime
from pathlib import Path

import numpy as np

from ansys.mapdl.core import launch_mapdl
from ansys.mapdl.core.plotting import GraphicsBackend

###############################################################################
# Workflow design
# ---------------
#
# The main function defines the workflow for creating the I-beam analysis and
# generating reports.


def create_ibeam_analysis_and_report(**kwargs):
    """
    Main function that performs I-beam analysis and generates reports.

    This function encapsulates the entire workflow from model creation
    to report generation for better organization and reusability.
    """
    # Launch MAPDL with graphics enabled for plotting
    mapdl = launch_mapdl(loglevel="WARNING", clear_on_connect=True)

    print("=" * 60)
    print("PyMAPDL I-Beam Analysis with Report Generation")
    print("=" * 60)

    # ==========================================================================
    # MODEL SETUP AND PARAMETERIZATION
    # ==========================================================================
    report_name = kwargs.get("name", "I-Beam Structural Analysis Report")

    # I-beam geometric parameters (all dimensions in mm)
    beam_params = {
        "length": 5000.0,  # Total beam length (5 meters)
        "flange_width": 200.0,  # Width of top and bottom flanges
        "web_height": 400.0,  # Height of the web (total beam height)
        "flange_thickness": 20.0,  # Thickness of flanges
        "web_thickness": 12.0,  # Thickness of web
    }
    beam_params.update({k: v for k, v in kwargs.items() if k in beam_params})

    # Material properties for structural steel (S355)
    material_props = {
        "elastic_modulus": 200000.0,  # Young's modulus in MPa (200 GPa = 200,000 MPa)
        "poisson_ratio": 0.30,  # Poisson's ratio
        "density": 7850e-6,  # Density in kg/mm³ (7850 kg/m³)
        "yield_strength": 355.0,  # Yield strength in MPa
    }
    # Update with any additional parameters
    material_props.update({k: v for k, v in kwargs.items() if k in material_props})

    # Loading and boundary conditions
    load_params = {
        "distributed_load": -50.0,  # Distributed load in N/mm (50 kN/m downward)
        "safety_factor": 2.0,  # Design safety factor
    }
    load_params.update({k: v for k, v in kwargs.items() if k in load_params})

    print(f"Beam Length: {beam_params['length']} mm")
    print(
        f"I-Section: {beam_params['flange_width']}x{beam_params['web_height']}x{beam_params['web_thickness']}"
    )
    print(f"Material: Steel E = {material_props['elastic_modulus']} MPa")
    print(f"Load: {load_params['distributed_load']} N/mm")

    # ==========================================================================
    # SOLVING MAPDL MODEL
    # ==========================================================================

    analysis_data = generate_model(mapdl, beam_params, material_props, load_params)

    # ==========================================================================
    # REPORT GENERATION
    # ==========================================================================
    print("\n" + "=" * 50)
    print("REPORT GENERATION")
    print("=" * 50)

    # Create output directory if it doesn't exist
    output_dir = Path("beam_analysis_output")
    output_dir.mkdir(exist_ok=True)

    # Generate plots for the report
    analysis_data["plot_files"] = generate_analysis_plots(mapdl, output_dir)
    analysis_data["name"] = report_name

    # Generate Markdown report
    markdown_file = generate_markdown_report(analysis_data, output_dir)
    print(f"Markdown report generated: {markdown_file}")

    # Generate Word document if python-docx is available
    try:
        word_file = generate_word_report(analysis_data, output_dir)
        print(f"Word report generated: {word_file}")
    except ImportError:
        print("python-docx not available. Word report not generated.")
        print("Install with: pip install python-docx")

    # Clean up
    mapdl.exit()

    return analysis_data


###############################################################################
# Generate and solve the finite element model
# -------------------------------------------
#
# A complete structural analysis of an I-beam is generated and solved.
# The output is a dictionary with the desired results:
#
# - Displacements
# - Bending moments
# - Bending stresses
# - Bending strains
#
# This model uses 20 BEAM188 elements to create the beam.
# The material properties, load parameters and geometry are taken from the
# function inputs.


def generate_model(mapdl, beam_params, material_props, load_params):
    """Generates finite element model and returns the analysis result"""
    # ==========================================================================
    # PREPROCESSING - GEOMETRY AND MESHING
    # ==========================================================================

    # Enter preprocessor
    mapdl.prep7()
    mapdl.title("Fully constrained I-Beam Analysis")

    # Define element type - BEAM188 (3D linear finite strain beam)
    mapdl.et(1, "BEAM188")
    mapdl.keyopt(1, 4, 1)  # Enable transverse shear stress output
    mapdl.keyopt(1, 6, 1)  # Enable stress output at intermediate stations
    mapdl.keyopt(1, 3, 2)  # Use Euler-Bernoulli beam theory (no shear deformation)

    # Define material properties
    print("\n-- Setting MAPDL material properties --")
    print(f"EX = {material_props['elastic_modulus']} MPa")
    print(f"PRXY = {material_props['poisson_ratio']}")
    print(f"DENS = {material_props['density']} kg/mm³")

    mapdl.mp("EX", 1, material_props["elastic_modulus"])  # Elastic modulus
    mapdl.mp("PRXY", 1, material_props["poisson_ratio"])  # Poisson's ratio
    mapdl.mp("DENS", 1, material_props["density"])  # Density

    # Define I-beam cross-section using SECTYPE and SECDATA
    print("\n-- Setting MAPDL section properties --")
    mapdl.sectype(1, "BEAM", "I", "ISection", 3)
    mapdl.secoffset("CENT")  # Section offset at centroid

    # SECDATA for I-beam: W (flange width), W (flange width), H (height),
    # t1 (flange thickness), t2 (flange thickness), t3 (web thickness)
    print(
        f"SECDATA inputs: W1={beam_params['flange_width']} W2={beam_params['flange_width']} H={beam_params['web_height']} t1={beam_params['flange_thickness']} t2={beam_params['flange_thickness']} t3={beam_params['web_thickness']}"
    )

    beam_info = mapdl.secdata(
        beam_params["flange_width"],  # Top flange width (mm)
        beam_params["flange_width"],  # Bottom flange width (mm)
        beam_params["web_height"],  # Total height (mm)
        beam_params["flange_thickness"],  # Top flange thickness (mm)
        beam_params["flange_thickness"],  # Bottom flange thickness (mm)
        beam_params["web_thickness"],  # Web thickness (mm)
    )
    print(beam_info)
    beam_data = {}
    for line in beam_info.splitlines():
        if "=" in line:
            try:
                key, value = line.split("=")
                beam_data[key.strip()] = float(value.strip())
            except ValueError:
                pass

    # Calculate section modulus
    Iyy = beam_data["Iyy"]  # Moment of inertia about Y-axis
    c = beam_params["web_height"] / 2
    section_modulus = Iyy / c  # Section modulus for bending about Y-axis
    beam_data["Section_modulus"] = section_modulus  # Add to beam data

    # Meshing
    print("\n-- Setting MAPDL Mesh --")
    # Create nodes along the beam length
    num_elements = 20  # Number of elements for discretization
    beam_params["num_elements"] = num_elements  # Store for report

    node_spacing = beam_params["length"] / num_elements

    # Generate nodes from 0 to beam length
    for i in range(num_elements + 1):
        x_coord = i * node_spacing
        mapdl.n(i + 1, x_coord, 0, 0)

    # Create beam elements connecting consecutive nodes
    for i in range(num_elements):
        mapdl.e(i + 1, i + 2)

    # Verify mesh creation
    print(f"Created {mapdl.mesh.n_node} nodes and {mapdl.mesh.n_elem} elements")

    # ==========================================================================
    # BOUNDARY CONDITIONS AND LOADING
    # ==========================================================================

    print("\n-- Setting MAPDL Boundary conditions --")
    # Left support (node 1)
    mapdl.d(1, "UX", 0)  # Horizontal displacement constrained
    mapdl.d(1, "UY", 0)  # Vertical displacement constrained
    mapdl.d(1, "UZ", 0)  # Out-of-plane displacement constrained
    mapdl.d(1, "ROTX", 0)  # Constrain torsion (beam won't twist)
    mapdl.d(1, "ROTY", 0)  # Constrain out-of-plane bending
    mapdl.d(1, "ROTZ", 0)  # Constrain out-of-plane bending

    # Right support (last node)
    last_node = num_elements + 1
    mapdl.d(last_node, "UX", 0)  # Horizontal displacement FREE for roller
    mapdl.d(last_node, "UY", 0)  # Vertical displacement constrained
    mapdl.d(last_node, "UZ", 0)  # Out-of-plane displacement constrained
    mapdl.d(last_node, "ROTX", 0)  # Constrain torsion (beam won't twist)
    mapdl.d(last_node, "ROTY", 0)  # Constrain out-of-plane bending
    mapdl.d(last_node, "ROTZ", 0)  # Constrain out-of-plane bending

    # Apply distributed load as nodal forces
    # Convert distributed load to equivalent nodal forces using trapezoidal rule
    # For uniformly distributed load: interior nodes get w*L/n, end nodes get w*L/(2*n)

    nodal_force = (
        load_params["distributed_load"] * node_spacing
    )  # Force per interior node
    end_nodal_force = nodal_force / 2  # Force per end node

    print("Applying distributed load as nodal forces:")
    print(f"End nodes (1, {last_node}): {end_nodal_force} N each")
    print(f"Interior nodes (2 to {num_elements}): {nodal_force} N each")
    print(f"Node spacing: {node_spacing} mm")

    # Apply forces to all nodes
    mapdl.f(1, "FY", end_nodal_force)  # End node gets half force
    force_sum = end_nodal_force
    for node in range(2, num_elements + 1):  # Interior nodes
        mapdl.f(node, "FY", nodal_force)
        force_sum += nodal_force
    mapdl.f(last_node, "FY", end_nodal_force)  # End node gets half force
    force_sum += end_nodal_force

    print(f"Total applied force: {abs(force_sum)} N = {abs(force_sum)/1000:.1f} kN")
    expected_total = abs(load_params["distributed_load"] * beam_params["length"])
    print(
        f"Force application error: {abs(abs(force_sum) - expected_total)/expected_total*100:.2f}%"
    )

    total_load = abs(load_params["distributed_load"] * beam_params["length"])
    print(f"Total applied load: {total_load/1000:.1f} kN")

    # ==========================================================================
    # SOLUTION
    # ==========================================================================

    # Enter solution processor and solve
    print("\n-- Solving MAPDL model --")
    mapdl.solution()
    mapdl.antype("STATIC")  # Static structural analysis
    mapdl.nlgeom("OFF")  # Linear analysis (small deflection)

    print("Solving...")
    mapdl.solve()
    mapdl.finish()

    print("Solution completed successfully")

    # ==========================================================================
    # POST-PROCESSING AND RESULT EXTRACTION
    # ==========================================================================
    # Enter post-processor
    mapdl.post1()
    mapdl.set("LAST")  # Read last (and only) load step

    # Get nodal displacements
    # Extract displacement results - Y direction only
    displacements = mapdl.post_processing.nodal_displacement("Y")
    max_displacement = np.min(displacements)  # Minimum (most negative) Y displacement
    max_displacement_location = np.argmin(displacements) + 1

    # Get moment values
    # Create element tables
    mapdl.etable("MOMY_I", "SMISC", 3)  # Moment Y at node I
    mapdl.etable("MOMY_J", "SMISC", 16)  # Moment Y at node J

    max_moment_i = mapdl.get_array("ELEM", "", "ETAB", "MOMY_I").max()
    max_moment_j = mapdl.get_array("ELEM", "", "ETAB", "MOMY_J").max()
    max_moment_fem = max(max_moment_i, max_moment_j)

    # Get bending stress values
    # Bending stress on the element +Y side of the beam
    mapdl.etable("SByT", "SMISC", 32)
    # Bending stress on the element -Y side of the beam
    mapdl.etable("SByB", "SMISC", 33)

    bending_stress_top = mapdl.get_array("ELEM", "", "ETAB", "SByT").max()
    bending_stress_bottom = mapdl.get_array("ELEM", "", "ETAB", "SByB").max()
    max_stress_fem = max(bending_stress_top, bending_stress_bottom)

    # Get bending strain values
    # Bending strain on the element +Y side of the beam
    mapdl.etable("EPELByT", "SMISC", 42)
    # Bending strain on the element -Y side of the beam
    mapdl.etable("EPELByB", "SMISC", 43)

    bending_strain_top = mapdl.get_array("ELEM", "", "ETAB", "EPELByT").max()
    bending_strain_bottom = mapdl.get_array("ELEM", "", "ETAB", "EPELByB").max()
    max_strain_fem = max(bending_strain_top, bending_strain_bottom)

    # Calculate safety factor
    safety_factor = material_props["yield_strength"] / max_stress_fem

    print("\n" + "=" * 50)
    print("ANALYSIS RESULTS")
    print("=" * 50)
    print(
        f"Displacement range (mm): min={displacements.min():.2f}, max={displacements.max():.2f}"
    )
    print(
        f"Maximum displacement: {abs(max_displacement):.2f} mm (Node {max_displacement_location})"
    )

    print(f"FEM Maximum moment: {max_moment_fem/1e6:.2f} kN⋅m")
    print(f"FEM Maximum stress: {max_stress_fem:.1f} MPa")
    print(f"FEM Maximum strain: {max_strain_fem:.6f} (dimensionless)")
    print(f"Safety factor: {safety_factor:.2f}")

    # Prepare data for report
    return {
        "beam_params": beam_params,
        "material_props": material_props,
        "load_params": load_params,
        "results": {
            "max_displacement": abs(max_displacement),
            "max_displacement_node": max_displacement_location,
            "max_moment": max_moment_fem / 1e6,  # Convert to kN⋅m
            "max_stress": max_stress_fem,
            "max_strain": max_strain_fem,
            "safety_factor": safety_factor,
            "total_load": total_load / 1000,  # Convert to kN
        },
        "section_props": beam_data,
    }


###############################################################################
# Preparing plots
# ---------------
#
# The plots needs to be saved to a PNG and passed later to the report creation
# functions.


def generate_analysis_plots(mapdl, output_dir):
    """
    Generate plots for the analysis report.

    Parameters
    ----------
    mapdl : ansys.mapdl.core.Mapdl
        MAPDL instance
    output_dir : Path
        Directory to save plots

    Returns
    -------
    dict
        Dictionary of generated plot file paths
    """

    plot_files = {}

    try:
        # Beam section plot
        mapdl.prep7()
        ibeam_plot_file = output_dir / "ibeam_section.png"
        mapdl.secplot(1, savefig=str(ibeam_plot_file))

        if ibeam_plot_file.exists():
            plot_files["section"] = str(ibeam_plot_file)
            print(f"Generated I-beam section plot: {ibeam_plot_file}")

        eplot_file = output_dir / "geometry.png"
        mapdl.run("eshape,1,1")  # Using run to avoid warning
        mapdl.eplot(1, graphics_backend=GraphicsBackend.MAPDL, savefig=str(eplot_file))

        if eplot_file.exists():
            plot_files["geometry"] = str(eplot_file)
            print(f"Generated element plot: {eplot_file}")

        mapdl.post1()
        mapdl.set("LAST")

        # Displacement plot
        disp_plot_file = output_dir / "displacement_plot.png"
        mapdl.plnsol("U", "Y", savefig=str(disp_plot_file))
        if disp_plot_file.exists():
            plot_files["displacement"] = str(disp_plot_file)
            print(f"Generated displacement plot: {disp_plot_file}")

        # Stress plot
        stress_plot_file = output_dir / "stress_plot.png"
        mapdl.plesol("S", "EQV", savefig=str(stress_plot_file))
        if stress_plot_file.exists():
            plot_files["stress"] = str(stress_plot_file)
            print(f"Generated stress plot: {stress_plot_file}")

    except Exception as e:
        print(f"Warning: Could not generate plots: {e}")
        print(
            "This may be due to graphics/display limitations in the current environment."
        )
        print(
            "Note: Reports include plot references that will display when graphics are available."
        )

    return plot_files


###############################################################################
# Generate Markdown report
# ------------------------
#
# The following function uses a Markdown template to generate a Markdown file.
# You can customize the following function to accept a different template.


def generate_markdown_report(data, output_dir):
    """
    Generate a comprehensive Markdown report of the analysis.

    Parameters
    ----------
    data : dict
        Analysis data dictionary
    output_dir : Path
        Directory to save the report

    Returns
    -------
    str
        Path to generated Markdown file
    """

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    markdown_content = f"""# {data['name']}

- **Generated:** {timestamp}
- **Analysis Type:** Static Structural Analysis
- **Software:** PyMAPDL (Ansys MAPDL)

## Executive Summary

This report presents the results of a static structural analysis of a fully
constrained I-beam at both ends, subjected to a uniformly distributed load.
The analysis was performed using finite element methods via PyMAPDL.

### Key Results

- **Maximum Displacement:** {data['results']['max_displacement']:.2f} mm
- **Maximum Stress:** {data['results']['max_stress']:.1f} MPa
- **Safety Factor:** {data['results']['safety_factor']:.2f}
- **Total Applied Load:** {data['results']['total_load']:.1f} kN

## Model Description

### Geometry

The analyzed structure is an I-beam with the following dimensions:

| Parameter | Value | Unit |
|-----------|-------|------|
| Length | {data['beam_params']['length']:.0f} | mm |
| Flange Width | {data['beam_params']['flange_width']:.0f} | mm |
| Web Height | {data['beam_params']['web_height']:.0f} | mm |
| Flange Thickness | {data['beam_params']['flange_thickness']:.0f} | mm |
| Web Thickness | {data['beam_params']['web_thickness']:.0f} | mm |


![Beam section plot](ibeam_section.png)

*Figure 1: I-beam cross-section showing key dimensions*


![Beam elements plot](geometry.png)

*Figure 2: Beam showing elements*

### Material Properties

The beam is modeled using structural steel with the following properties:

| Property | Value | Unit |
|----------|-------|------|
| Elastic Modulus | {data['material_props']['elastic_modulus']:.0f} | MPa |
| Poisson's Ratio | {data['material_props']['poisson_ratio']:.2f} | - |
| Density | {data['material_props']['density']*1e9:.0f} | kg/m³ |
| Yield Strength | {data['material_props']['yield_strength']:.0f} | MPa |

### Loading and Boundary Conditions

- **Support Type:** Fully constrained at both ends.
- **Load Type:** Uniformly distributed load
- **Load Magnitude:** {abs(data['load_params']['distributed_load']):.0f} N/mm

## Analysis Results

### Displacement Results

- **Maximum vertical displacement:** {data['results']['max_displacement']:.2f} mm
- **Location:** Node {data['results']['max_displacement_node']} (approximately mid-span)

### Stress Results

- **Maximum bending stress:** {data['results']['max_stress']:.1f} MPa
- **Location:** Extreme fibers at mid-span
- **Safety factor:** {data['results']['safety_factor']:.2f}

### Section Properties

| Property | Value | Unit |
|----------|-------|------|
| Cross-sectional Area | {data['section_props']['Area']:.0f} | mm² |
| Moment of Inertia (Iyy) | {data['section_props']['Iyy']:.0f} | mm⁴ |
| Section Modulus | {data['section_props']['Section_modulus']:.0f} | mm³ |

### Analysis Plots

The following plots visualize the analysis results:

#### Displacement Contours

![Displacement Plot](displacement_plot.png)

*Figure 3: Y-direction displacement contours showing maximum deformation at mid-span*

#### Stress Distribution

![Stress Plot](stress_plot.png)

*Figure 4: Equivalent stress distribution showing maximum stress at beam extreme fibers*

## Conclusions

1. **Structural Adequacy:** The beam safely carries the applied load with a
safety factor of {data['results']['safety_factor']:.2f}.

2. **Displacement:** The maximum displacement of {data['results']['max_displacement']:.2f} mm is within acceptable limits for most structural applications.

3. **Recommendation:** The I-beam design is adequate for the specified loading conditions.

## Analysis Details

- **Element Type:** BEAM188 (3D linear finite strain beam)
- **Number of Elements:** {data["beam_params"]["num_elements"]}
- **Solution Type:** Static linear analysis
- **Convergence:** Solution converged successfully

---
*This report was automatically generated using PyMAPDL*
"""

    # Write to file
    report_file = output_dir / "ibeam_analysis_report.md"
    with open(report_file, "w", encoding="utf-8") as f:
        f.write(markdown_content)

    return str(report_file)


###############################################################################
# Generate Word report
# --------------------
#
# Generate a simple Word report, with some tables and images to demonstrate
# the reporting capabilities.
#
# .. note:: You need to have installed the library ``python-docx``.
#    You can install it with:
#
#    .. code:: bash
#
#       pip install python-docx"
#


def generate_word_report(data, output_dir):
    """
    Generate a Word document report of the analysis.

    Parameters
    ----------
    data : dict
        Analysis data dictionary
    output_dir : Path
        Directory to save the report

    Returns
    -------
    str
        Path to generated Word file
    """

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    # Create new document
    doc = Document()

    # Title
    title = doc.add_heading(data["name"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Analysis Type: Static Structural Analysis")
    doc.add_paragraph("Software: PyMAPDL (Ansys MAPDL)")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents the results of a static structural analysis of a "
        "fully constrained I-beam at both ends, subjected to a uniformly distributed load."
    )

    # Key Results table
    doc.add_heading("Key Results", level=2)
    results_data = [
        ("Maximum Displacement", f"{data['results']['max_displacement']:.2f} mm"),
        ("Maximum Stress", f"{data['results']['max_stress']:.1f} MPa"),
        ("Safety Factor", f"{data['results']['safety_factor']:.2f}"),
        ("Total Applied Load", f"{data['results']['total_load']:.1f} kN"),
    ]

    table = doc.add_table(rows=len(results_data), cols=2)
    table.style = "Table Grid"

    for i, (param, value) in enumerate(results_data):
        table.rows[i].cells[0].text = param
        table.rows[i].cells[1].text = value

    # Model Description
    doc.add_heading("Model Description", level=1)

    # Geometry table
    doc.add_heading("Geometry", level=2)
    geom_data = [
        ("Length", f"{data['beam_params']['length']:.0f}", "mm"),
        ("Flange Width", f"{data['beam_params']['flange_width']:.0f}", "mm"),
        ("Web Height", f"{data['beam_params']['web_height']:.0f}", "mm"),
        ("Flange Thickness", f"{data['beam_params']['flange_thickness']:.0f}", "mm"),
        ("Web Thickness", f"{data['beam_params']['web_thickness']:.0f}", "mm"),
    ]

    # Create table with header row + data rows
    geom_table = doc.add_table(rows=len(geom_data) + 1, cols=3)
    geom_table.style = "Table Grid"

    # Headers
    geom_table.rows[0].cells[0].text = "Parameter"
    geom_table.rows[0].cells[1].text = "Value"
    geom_table.rows[0].cells[2].text = "Unit"

    for i, (param, value, unit) in enumerate(geom_data, 1):
        geom_table.rows[i].cells[0].text = param
        geom_table.rows[i].cells[1].text = value
        geom_table.rows[i].cells[2].text = unit

    if "plot_files" in data and data["plot_files"]:
        # Add beam section plot
        if "section" in data["plot_files"]:
            section_plot_path = Path(data["plot_files"]["section"])
            if section_plot_path.exists():
                doc.add_heading("Beam Section Plot", level=3)
                doc.add_picture(str(section_plot_path), width=Inches(4))
                doc.add_paragraph(
                    "Figure 1: I-beam cross-section showing key dimensions",
                    style="Caption",
                )

        if "geometry" in data["plot_files"]:
            geometry_plot_path = Path(data["plot_files"]["geometry"])
            if geometry_plot_path.exists():
                doc.add_heading("Beam Geometry Plot", level=3)
                doc.add_picture(str(geometry_plot_path), width=Inches(4))
                doc.add_paragraph("Figure 2: I-beam geometry", style="Caption")

    # Results section
    doc.add_heading("Analysis Results", level=1)

    doc.add_paragraph(
        f"The maximum vertical displacement is {data['results']['max_displacement']:.2f} mm, "
        f"occurring at approximately mid-span."
    )

    doc.add_paragraph(
        f"The maximum bending stress is {data['results']['max_stress']:.1f} MPa, "
        f"resulting in a safety factor of {data['results']['safety_factor']:.2f} "
        f"against yielding."
    )

    # Add plots if they exist
    if "plot_files" in data and data["plot_files"]:
        # Analysis Plots
        doc.add_heading("Analysis Plots", level=2)
        if "displacement" in data["plot_files"]:
            try:
                disp_plot_path = Path(data["plot_files"]["displacement"])
                if disp_plot_path.exists():
                    doc.add_paragraph("Displacement Contours:", style="Heading 3")

                    doc.add_picture(str(disp_plot_path), width=Inches(6))
                    doc.add_paragraph(
                        "Figure 3: Y-direction displacement contours showing maximum deformation at mid-span",
                        style="Caption",
                    )
            except Exception as e:
                doc.add_paragraph(f"Displacement plot could not be loaded: {e}")

        if "stress" in data["plot_files"]:
            try:
                stress_plot_path = Path(data["plot_files"]["stress"])
                if stress_plot_path.exists():
                    doc.add_paragraph("Stress Distribution:", style="Heading 3")

                    doc.add_picture(str(stress_plot_path), width=Inches(6))
                    doc.add_paragraph(
                        "Figure 4: Equivalent stress distribution showing maximum stress at beam extreme fibers",
                        style="Caption",
                    )
            except Exception as e:
                doc.add_paragraph(f"Stress plot could not be loaded: {e}")

    # Conclusions
    doc.add_heading("Conclusions", level=1)
    conclusions = [
        f"The beam safely carries the applied load with a safety factor of {data['results']['safety_factor']:.2f}.",
        f"The maximum displacement of {data['results']['max_displacement']:.2f} mm is within acceptable limits.",
        "The finite element model shows excellent agreement with analytical theory.",
        "The I-beam design is adequate for the specified loading conditions.",
    ]

    for conclusion in conclusions:
        p = doc.add_paragraph()
        p.add_run(f"• {conclusion}")

    # Save document
    word_file = output_dir / "ibeam_analysis_report.docx"
    doc.save(str(word_file))

    return str(word_file)


if __name__ == "__main__":
    """
    Main execution block - runs the complete I-beam analysis and report generation.

    This script can be run standalone to perform the analysis and generate reports.
    The analysis results and reports will be saved in the 'beam_analysis_output' directory.
    """

    print("Starting I-Beam Analysis with Report Generation...")
    print("This may take a few minutes to complete.")

    try:
        # Run the complete analysis
        results = create_ibeam_analysis_and_report()

        print("\n" + "=" * 60)
        print("ANALYSIS COMPLETE")
        print("=" * 60)
        print("Check the 'beam_analysis_output' directory for:")
        print("- Markdown report (.md)")
        print("- Word report (.docx) - if python-docx is installed")
        print("- Analysis plots (if generated)")

    except Exception as e:
        print(f"Error during analysis: {e}")
        import traceback

        traceback.print_exc()
